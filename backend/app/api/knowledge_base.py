"""
Knowledge Base API - PDF Upload, Processing, and Management
Supports: Multiple PDF uploads, Full text extraction, Intelligent chunking
"""

from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, BackgroundTasks
from fastapi.responses import JSONResponse
from typing import List, Dict, Any, Optional
from collections import defaultdict
from pydantic import BaseModel
import os
import logging
from pathlib import Path
import shutil
from datetime import datetime

# Lazy import to avoid scipy/sklearn initialization errors with Python 3.14
# These will be imported only when actually needed
# from app.services.enhanced_knowledge_base import get_knowledge_base
# from app.services.local_vector_kb import get_local_knowledge_base
from app.services.vector_knowledge_base import get_vector_knowledge_base
from app.api.auth_new import get_current_user
from app.models import User, KnowledgeDocument, UserRole
from app.database import get_db
from sqlalchemy.orm import Session
import hashlib
import hashlib
import uuid
import os
import openai
from app.services.online_knowledge_service import OnlineKnowledgeService

logger = logging.getLogger(__name__)

router = APIRouter(tags=["knowledge-base"])

# ==================== Access Control ====================

def require_kb_management_role(current_user: User = Depends(get_current_user)) -> User:
    """
    Dependency to ensure user has KB management permissions (Admin or Doctor only).
    Raises 403 if user is not authorized.
    """
    if current_user.role not in [UserRole.ADMIN, UserRole.DOCTOR]:
        logger.warning(f"KB access denied for user {current_user.email} with role {current_user.role}")
        raise HTTPException(
            status_code=403,
            detail=f"Access denied. Knowledge Base management requires Admin or Doctor role. Your role: {current_user.role.value}"
        )
    return current_user

def require_admin_role(current_user: User = Depends(get_current_user)) -> User:
    """
    Dependency to ensure user is Admin only.
    Raises 403 if user is not an admin.
    """
    if current_user.role != UserRole.ADMIN:
        logger.warning(f"Admin access denied for user {current_user.email} with role {current_user.role}")
        raise HTTPException(
            status_code=403,
            detail=f"Access denied. This operation requires Admin role. Your role: {current_user.role.value}"
        )
    return current_user

@router.get("/test")
async def test_kb():
    """Test endpoint to verify KB router is working"""
    logger.info("TEST ENDPOINT CALLED")
    return {"status": "test", "message": "KB router is working!"}

# Constants
MAX_FILE_SIZE = 1024 * 1024 * 1024  # 1GB per file (increased for large medical textbooks)
MAX_TOTAL_SIZE = 5 * 1024 * 1024 * 1024  # 5GB total
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".doc", ".docx"}
UPLOAD_DIR = Path("data/knowledge_base/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

def infer_year_from_name(name: str) -> Optional[int]:
    import re
    current_year = datetime.now().year + 1
    matches = re.findall(r"(19\d{2}|20\d{2})", name)
    for match in matches:
        year = int(match)
        if 1950 <= year <= current_year:
            return year
    return None

# Import large PDF processor
from app.services.large_pdf_processor import get_large_pdf_processor
from app.services.error_corrector import get_error_corrector, with_auto_correction
from app.services.pdf_ocr_processor import get_pdf_ocr_processor


class PDFProcessingRequest(BaseModel):
    use_full_content: bool = False  # False = intelligent chunking (DEFAULT), True = full PDF
    chunk_size: int = 2000  # Characters per chunk (larger = fewer chunks = faster)
    overlap: int = 50  # Overlap between chunks (smaller = fewer chunks)
    extract_tables: bool = False  # Set to True if you need tables (slower)
    extract_images: bool = True  # Extract images with metadata (DEFAULT: True)
    ocr_enabled: bool = True  # Enable OCR for scanned PDFs (DEFAULT: True)


class SearchRequest(BaseModel):
    query: str
    top_k: int = 5
    min_score: float = 0.0
    search_mode: str = "hybrid"  # local|openai|hybrid
    alpha: float = 0.6  # hybrid weight
    filters: Optional[Dict[str, Any]] = None
    allow_fallback: bool = True
    synthesize_answer: bool = False  # New: Generate consolidated answer


class EvaluationItem(BaseModel):
    query: str
    expected_terms: List[str]
    top_k: int = 5


class EvaluationRequest(BaseModel):
    runs: List[EvaluationItem]


@router.post("/upload")
async def upload_pdfs(
    files: List[UploadFile] = File(...),
    use_full_content: bool = False,  # Default to chunking for semantic search
    chunk_size: int = 1000,
    category: str = "medical_textbook",  # Default category
    extract_images: bool = True,  # Extract and index images
    ocr_enabled: bool = True,  # Enable OCR for scanned PDFs
    current_user: User = Depends(require_kb_management_role),
    db: Session = Depends(get_db)
):
    """
    Upload multiple PDF files to knowledge base.
    **Requires Admin or Doctor role.**
    
    - **files**: List of PDF files (up to 200MB each, 1GB total)
    - **use_full_content**: If False (default), intelligent chunking; if True, full PDF as single document
    - **chunk_size**: Size of text chunks when use_full_content=False
    - **category**: Document category (medical_textbook, clinical_guidelines, research_paper, etc.)
    - **extract_images**: Extract and save images with metadata (default: True)
    - **ocr_enabled**: Enable OCR for scanned/image-based PDFs (default: True)
    """
    if not files:
        raise HTTPException(status_code=400, detail="No files provided")
    
    # Validate file count
    if len(files) > 20:
        raise HTTPException(status_code=400, detail="Maximum 20 files allowed per upload")

    # Stream files to disk while computing size and hash to avoid double reads on large uploads
    total_size = 0
    prepared_files = []
    for file in files:
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"File {file.filename} has unsupported type. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
            )

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = f"{timestamp}_{file.filename}"
        file_path = UPLOAD_DIR / safe_filename

        hasher = hashlib.sha256()
        file_size = 0
        try:
            with open(file_path, "wb") as f:
                while chunk := await file.read(4 * 1024 * 1024):  # 4MB chunks for faster disk writes
                    file_size += len(chunk)
                    if file_size > MAX_FILE_SIZE:
                        raise HTTPException(
                            status_code=400,
                            detail=f"File {file.filename} exceeds maximum size of {MAX_FILE_SIZE / 1024 / 1024:.0f}MB"
                        )
                    hasher.update(chunk)
                    f.write(chunk)
        except HTTPException:
            if file_path.exists():
                file_path.unlink()
            raise
        except Exception as e:
            if file_path.exists():
                file_path.unlink()
            raise HTTPException(status_code=500, detail=f"Failed to save {file.filename}: {str(e)}")

        if file_size == 0:
            if file_path.exists():
                file_path.unlink()
            raise HTTPException(status_code=400, detail=f"File {file.filename} appears to be empty")

        total_size += file_size
        if total_size > MAX_TOTAL_SIZE:
            file_path.unlink(missing_ok=True)
            raise HTTPException(
                status_code=400,
                detail=f"Total upload size ({total_size / 1024 / 1024:.2f}MB) exceeds maximum of {MAX_TOTAL_SIZE / 1024 / 1024 / 1024:.1f}GB"
            )

        prepared_files.append({
            "file": file,
            "file_ext": file_ext,
            "file_path": file_path,
            "file_size": file_size,
            "file_hash": hasher.hexdigest(),
            "safe_filename": safe_filename
        })

    # Process files
    results = []
    # Lazy import to avoid scipy/sklearn initialization errors
    from app.services.enhanced_knowledge_base import get_knowledge_base
    knowledge_base = get_knowledge_base()
    
    for prepared in prepared_files:
        file = prepared["file"]
        file_ext = prepared["file_ext"]
        file_path = prepared["file_path"]
        file_size = prepared["file_size"]
        file_hash = prepared["file_hash"]
        try:
            logger.info(f"Processing file: {file.filename}")
            
            # Check if document already exists
            existing_doc = db.query(KnowledgeDocument).filter(
                KnowledgeDocument.file_hash == file_hash
            ).first()
            
            if existing_doc:
                logger.info(f"Document {file.filename} already exists (hash: {file_hash})")
                results.append({
                    "filename": file.filename,
                    "status": "skipped",
                    "reason": "Document already uploaded",
                    "existing_document_id": existing_doc.document_id,
                    "uploaded_at": existing_doc.uploaded_at.isoformat()
                })
                continue
            
            # Extract text based on file type
            file_ext = Path(file.filename).suffix.lower()
            extracted_images = []
            extraction_method = "text"
            
            # Generate unique document ID early for OCR usage
            doc_uuid = str(uuid.uuid4())
            
            if file_ext == ".pdf":
                # Use enhanced OCR processor
                ocr_processor = get_pdf_ocr_processor()
                pdf_result = ocr_processor.extract_pdf_with_images(
                    file_path,
                    extract_images=extract_images,
                    use_ocr=ocr_enabled,
                    document_id=doc_uuid
                )
                text_content = pdf_result['text']
                extracted_images = pdf_result.get('images', [])
                extraction_method = pdf_result.get('method', 'text')
                
                logger.info(f"[EXTRACT] {file.filename}: {len(text_content)} chars, {len(extracted_images)} images, method={extraction_method}")
            elif file_ext == ".txt":
                text_content = file_path.read_bytes().decode('utf-8', errors='ignore')
            elif file_ext in [".doc", ".docx"]:
                text_content = await extract_word_text(file_path)
            else:
                text_content = ""
            
            if not text_content.strip():
                results.append({
                    "filename": file.filename,
                    "status": "error",
                    "error": "No text content extracted - PDF may be image-based or encrypted. Try using OCR or saving as text-based PDF.",
                    "chunks": 0,
                    "characters": 0,
                    "file_size_mb": file_size / 1024 / 1024
                })
                continue
            
            # Minimum character threshold (50 chars to filter out near-empty extractions)
            if len(text_content.strip()) < 50:
                results.append({
                    "filename": file.filename,
                    "status": "error",
                    "error": f"Insufficient text extracted ({len(text_content)} chars). PDF appears to be image-based. OCR required.",
                    "chunks": 0,
                    "characters": len(text_content),
                    "file_size_mb": file_size / 1024 / 1024
                })
                continue
            
            # Process content
            if use_full_content:
                # Add as single document
                doc_id = await add_to_knowledge_base(
                    knowledge_base,
                    text_content,
                    file.filename,
                    metadata={
                        "type": "full_document",
                        "uploaded_by": current_user.email,
                        "upload_date": datetime.now().isoformat(),
                        "document_uuid": doc_uuid,
                        "filename": file.filename,
                        "category": "medical_pdf",
                        "section": "full_document",
                        "year": infer_year_from_name(file.filename),
                        "outdated": False
                    }
                )
                
                # Save to database
                db_doc = KnowledgeDocument(
                    document_id=doc_uuid,
                    filename=file.filename,
                    file_path=str(file_path),
                    file_hash=file_hash,
                    file_size=file_size,
                    extension=file_ext,
                    text_length=len(text_content),
                    chunk_count=1,
                    is_indexed=True,
                    indexed_at=datetime.now(),
                    uploaded_by_id=current_user.id,
                    category=category,  # Use parameter
                    source=file.filename
                )
                db.add(db_doc)
                db.commit()
                
                results.append({
                    "filename": file.filename,
                    "status": "success",
                    "processing_mode": "full_content",
                    "document_id": doc_uuid,
                    "chunks": 1,
                    "characters": len(text_content),
                    "size_mb": file_size / 1024 / 1024,
                    "images_extracted": len(extracted_images),
                    "extraction_method": extraction_method,
                    "embedding_status": "queued_for_background_processing",
                    "info": f"Extracted via {extraction_method}. {len(extracted_images)} images saved. Embeddings queued for background processing."
                })
            else:
                # Intelligent chunking with reduced overlap for speed
                chunks = smart_chunk_text(text_content, chunk_size=chunk_size, overlap=50)
                
                chunk_ids = []
                for i, chunk in enumerate(chunks):
                    doc_id = await add_to_knowledge_base(
                        knowledge_base,
                        chunk,
                        f"{file.filename} (Part {i+1}/{len(chunks)})",
                        metadata={
                            "type": "chunk",
                            "chunk_index": i,
                            "total_chunks": len(chunks),
                            "uploaded_by": current_user.email,
                            "upload_date": datetime.now().isoformat(),
                            "document_uuid": doc_uuid,
                            "filename": file.filename,
                            "category": "medical_pdf",
                            "section": "body",
                            "year": infer_year_from_name(file.filename),
                            "outdated": False
                        }
                    )
                    chunk_ids.append(doc_id)
                
                db_doc = KnowledgeDocument(
                    document_id=doc_uuid,
                    filename=file.filename,
                    file_path=str(file_path),
                    file_hash=file_hash,
                    file_size=file_size,
                    extension=file_ext,
                    text_length=len(text_content),
                    chunk_count=len(chunks),
                    is_indexed=True,
                    indexed_at=datetime.now(),
                    uploaded_by_id=current_user.id,
                    category=category,  # Use parameter
                    source=file.filename
                )
                db.add(db_doc)
                db.commit()
                
                results.append({
                    "filename": file.filename,
                    "status": "success",
                    "processing_mode": "intelligent_chunking",
                    "document_id": doc_uuid,
                    "document_ids": chunk_ids,
                    "chunks": len(chunks),
                    "characters": len(text_content),
                    "avg_chunk_size": len(text_content) // len(chunks) if chunks else 0,
                    "size_mb": file_size / 1024 / 1024,
                    "images_extracted": len(extracted_images),
                    "extraction_method": extraction_method,
                    "embedding_status": "queued_for_background_processing",
                    "info": f"Extracted via {extraction_method}. {len(extracted_images)} images saved. {len(chunks)} chunks queued for embeddings."
                })
            
            logger.info(f"Successfully queued {file.filename}: {len(text_content)} characters")
            
        except Exception as e:
            error_type = type(e).__name__
            error_msg = str(e).strip() if str(e).strip() else "Unknown error"
            full_error = f"{error_type}: {error_msg}" if error_msg != "Unknown error" else f"{error_type} during PDF extraction/processing"
            
            logger.error(f"Error processing {file.filename}: {full_error}")
            import traceback
            logger.error(traceback.format_exc())
            
            results.append({
                "filename": file.filename,
                "status": "error",
                "error": full_error,
                "chunks": 0,
                "characters": 0,
                "file_size_mb": file_size / 1024 / 1024 if 'file_size' in locals() else 0
            })
    
    # Summary
    successful = sum(1 for r in results if r["status"] == "success")
    failed = len(results) - successful
    total_chunks = sum(r.get("chunks", 0) for r in results)
    
    return {
        "message": f"Processed {len(files)} files: {successful} successful, {failed} failed",
        "summary": {
            "total_files": len(files),
            "successful": successful,
            "failed": failed,
            "total_chunks_created": total_chunks,
            "total_size_mb": total_size / 1024 / 1024
        },
        "results": results,
        "status_check_endpoints": {
            "all_uploads": "/api/medical/knowledge/upload-status",
            "specific_upload": "/api/medical/knowledge/upload-status/{document_id}"
        }
    }


async def extract_pdf_text(file_path: Path, use_full_content: bool = True) -> str:
    """
    Extract text from PDF using PyMuPDF (fitz).
    Supports full extraction with formatting preservation.
    """
    try:
        import fitz  # PyMuPDF
        
        logger.info(f"Opening PDF: {file_path}")
        doc = fitz.open(str(file_path))
        page_count = len(doc)
        logger.info(f"PDF opened successfully: {page_count} pages")
        
        text_parts = []
        
        for page_num, page in enumerate(doc, 1):
            # Extract text with layout preservation
            text = page.get_text("text")
            
            if use_full_content:
                # Include page headers and formatting
                text_parts.append(f"\n--- Page {page_num} ---\n")
                text_parts.append(text)
            else:
                # Clean text for chunking
                text_parts.append(text)
            
            # Log progress every 50 pages
            if page_num % 50 == 0:
                logger.info(f"Processed {page_num}/{page_count} pages...")
        
        doc.close()
        full_text = "\n".join(text_parts)
        
        logger.info(f"[OK] Extracted {len(full_text)} characters from PDF with {page_count} pages")
        return full_text
        
    except ImportError as e:
        logger.warning(f"PyMuPDF not available: {e}, using fallback")
        # Fallback to basic text extraction
        return await extract_pdf_text_fallback(file_path)
    except Exception as e:
        logger.error(f"[ERROR] PDF extraction error for {file_path}: {type(e).__name__}: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return ""


async def extract_pdf_text_fallback(file_path: Path) -> str:
    """Fallback PDF extraction using PyPDF2"""
    try:
        import PyPDF2
        
        text_parts = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Fallback PDF extraction error: {e}")
        return ""


async def extract_word_text(file_path: Path) -> str:
    """Extract text from Word documents"""
    try:
        import docx
        
        doc = docx.Document(str(file_path))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    except ImportError:
        logger.warning("python-docx not available")
        return ""
    except Exception as e:
        logger.error(f"Word extraction error: {e}")
        return ""


def smart_chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """
    Intelligent text chunking that respects:
    - Paragraph boundaries
    - Sentence boundaries
    - Section headers
    """
    chunks = []
    
    # Split by paragraphs first
    paragraphs = text.split('\n\n')
    
    current_chunk = ""
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        # If adding this paragraph exceeds chunk_size
        if len(current_chunk) + len(para) > chunk_size:
            if current_chunk:
                chunks.append(current_chunk)
                # Add overlap from end of previous chunk
                overlap_text = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                current_chunk = overlap_text + "\n\n" + para
            else:
                # Paragraph itself is too long, split by sentences
                sentences = para.split('. ')
                for sentence in sentences:
                    if len(current_chunk) + len(sentence) > chunk_size:
                        if current_chunk:
                            chunks.append(current_chunk)
                            current_chunk = current_chunk[-overlap:] + ". " + sentence
                        else:
                            # Sentence too long, force split
                            chunks.append(sentence[:chunk_size])
                            current_chunk = sentence[chunk_size:]
                    else:
                        current_chunk += (". " if current_chunk else "") + sentence
        else:
            current_chunk += ("\n\n" if current_chunk else "") + para
    
    # Add remaining chunk
    if current_chunk:
        chunks.append(current_chunk)
    
    return chunks


async def add_to_knowledge_base(kb, text: str, source: str, metadata: Dict[str, Any]) -> str:
    """Add document to knowledge base with LOCAL EMBEDDING (NO API, NO COST, FAST)
    
    Returns immediately with doc ID - actual embedding happens async in background.
    This prevents blocking the HTTP response during large PDF processing.
    """
    try:
        # DEFERRED EMBEDDING: Don't block response - just queue for processing
        from app.services.local_vector_kb import get_local_knowledge_base
        local_kb = get_local_knowledge_base()
        
        logger.info(f"[QUEUED] Adding to LOCAL KB (deferred): {source} ({len(text)} chars)")
        
        # Prepare metadata with document_id for linking
        full_metadata = {
            "source": source,
            **metadata
        }
        # Enrich metadata for filtering/citations
        full_metadata.setdefault("filename", metadata.get("filename", source))
        full_metadata.setdefault("category", metadata.get("category", "medical_pdf"))
        full_metadata.setdefault("section", metadata.get("section", "general"))
        full_metadata.setdefault("outdated", metadata.get("outdated", False))
        inferred_year = infer_year_from_name(source)
        if inferred_year:
            full_metadata.setdefault("year", inferred_year)
        else:
            full_metadata.setdefault("year", metadata.get("year"))
        
        # Ensure document_id is available for reference links
        if 'document_uuid' in metadata:
            full_metadata['document_id'] = metadata['document_uuid']
        
        # ASYNC ADD: This queues the document without blocking
        # In production, this would use Celery/RQ/APScheduler
        # For now, we do a quick add without re-indexing
        import asyncio
        from pathlib import Path
        
        # Store text temporarily to process later
        temp_dir = Path("data/knowledge_base/pending")
        temp_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate a pending task ID
        import hashlib
        task_id = hashlib.md5(f"{source}{datetime.now().isoformat()}".encode()).hexdigest()[:16]
        pending_file = temp_dir / f"{task_id}.json"
        
        import json
        with open(pending_file, 'w') as f:
            json.dump({
                "task_id": task_id,
                "source": source,
                "text": text[:10000],  # Store first 10KB inline
                "text_length": len(text),
                "metadata": full_metadata,
                "created_at": datetime.now().isoformat(),
                "status": "pending"
            }, f)
        
        # Schedule async processing (non-blocking)
        asyncio.create_task(
            _process_document_async(local_kb, text, source, full_metadata, task_id)
        )
        
        logger.info(f"[QUEUED] {source} - scheduled for background processing (ID: {task_id})")
        return f"queued_{task_id}"
        
    except Exception as e:
        logger.error(f"Failed to queue KB document: {e}")
        import hashlib
        doc_id = hashlib.md5(f"{source}{text[:100]}".encode()).hexdigest()[:12]
        return doc_id


async def _process_document_async(kb, text: str, source: str, metadata: Dict[str, Any], task_id: str):
    """Background task to process document embeddings (does not block HTTP response)"""
    try:
        logger.info(f"[BACKGROUND] Processing embeddings for: {source} (task: {task_id})")
        
        chunks_added = kb.add_document(
            content=text,
            metadata=metadata,
            chunk_size=2000,
            chunk_overlap=50
        )
        
        # Update pending status
        from pathlib import Path
        pending_file = Path("data/knowledge_base/pending") / f"{task_id}.json"
        if pending_file.exists():
            import json
            with open(pending_file, 'r') as f:
                data = json.load(f)
            data['status'] = 'complete'
            data['chunks_added'] = chunks_added
            with open(pending_file, 'w') as f:
                json.dump(data, f)
        
        logger.info(f"[OK] Background embedding complete: {source} - {chunks_added} chunks (task: {task_id})")
    except Exception as e:
        logger.error(f"[ERROR] Background processing failed for {source}: {e}")
        # Log error but don't crash


@router.get("/statistics")
async def get_statistics():
    """Get knowledge base statistics (aggregated).

    Returns a structure tailored for the frontend KnowledgeBase page:
    - total_documents: number of indexed documents in database
    - total_chunks: number of indexed chunks (aggregated from documents)
    - knowledge_level: derived from chunk count
    - search_mode: current KB mode
    - pdf_sources: known PDF files and whether they are indexed
    - categories: unique document categories
    """
    logger.info("==== STATISTICS ENDPOINT CALLED ====")
    try:
        # Local vector KB (sentence-transformers + FAISS, no API cost)
        logger.info("About to call get_local_knowledge_base()...")
        from app.services.local_vector_kb import get_local_knowledge_base
        local_kb = get_local_knowledge_base()
        logger.info("Successfully got local KB instance")
        local_stats = local_kb.get_statistics()
        logger.info("Successfully got local KB statistics")

        # Query database directly for document count and chunks
        db_gen = get_db()
        db = next(db_gen)
        try:
            # Get document count and total chunks from database
            db_docs = db.query(KnowledgeDocument).all()
            db_doc_count = len(db_docs)
            db_total_chunks = sum(doc.chunk_count or 0 for doc in db_docs)
            
            # Get unique categories
            db_categories = set()
            for doc in db_docs:
                if doc.category and doc.category.strip():
                    db_categories.add(doc.category.strip())
            
            logger.info(f"Database query: {db_doc_count} documents, {db_total_chunks} chunks, {len(db_categories)} categories")
        finally:
            db.close()
        
        # Enhanced KB (medical database, optional re-ranking) - keep for capabilities
        from app.services.enhanced_knowledge_base import get_knowledge_base
        enhanced_kb = get_knowledge_base()
        enhanced_stats = enhanced_kb.get_statistics()
        
        # Use database count as primary source
        enhanced_doc_count = db_doc_count

        # Uploads directory (files uploaded via API/UI)
        upload_files = [f for f in UPLOAD_DIR.glob("*") if f.is_file()]
        total_upload_size = sum(f.stat().st_size for f in upload_files)

        # Medical books directory (pre-seeded PDFs)
        MEDICAL_BOOKS_DIR = Path("data/medical_books")
        medical_books_files: List[Path] = []
        if MEDICAL_BOOKS_DIR.exists():
            medical_books_files = [
                f for f in MEDICAL_BOOKS_DIR.glob("**/*")
                if f.is_file() and f.suffix.lower() in {".pdf", ".txt"}
            ]

        # Determine which sources are indexed by looking at metadata 'source'
        indexed_sources = set()
        try:
            for doc_meta in getattr(local_kb, 'documents', []):
                src = str(doc_meta.get('source') or '').strip()
                if src:
                    # Normalize base filename (strip chunk labels like "(Part x/y)")
                    base = src.split(' (Part')[0].strip()
                    indexed_sources.add(base)
        except Exception:
            pass

        def _file_entry(p: Path) -> Dict[str, Any]:
            name = p.name
            status = 'indexed' if name in indexed_sources else 'available'
            return {
                'name': name,
                'size_mb': round(p.stat().st_size / 1024 / 1024, 2),
                'status': status
            }

        pdf_sources: List[Dict[str, Any]] = []
        # Include medical_books first, then uploads
        pdf_sources.extend([_file_entry(p) for p in medical_books_files[:100]])  # cap list for payload size
        pdf_sources.extend([_file_entry(p) for p in upload_files[:100]])

        # Compute knowledge level from chunks
        total_chunks = int(local_stats.get('total_chunks', 0))
        if total_chunks >= 20000:
            knowledge_level = 'expert'
        elif total_chunks >= 10000:
            knowledge_level = 'advanced'
        elif total_chunks >= 1000:
            knowledge_level = 'intermediate'
        elif total_chunks > 0:
            knowledge_level = 'basic'
        else:
            knowledge_level = 'unknown'

        # Determine search mode
        search_mode = 'local-semantic' if local_stats.get('faiss_available') else 'simple'

        # Get processing queue status
        from app.models import DocumentProcessingStatus
        db_gen = get_db()
        db_status = next(db_gen)
        try:
            queued_count = db_status.query(DocumentProcessingStatus).filter(
                DocumentProcessingStatus.status == "queued"
            ).count()
            processing_count = db_status.query(DocumentProcessingStatus).filter(
                DocumentProcessingStatus.status == "processing"
            ).count()
            completed_count = db_status.query(DocumentProcessingStatus).filter(
                DocumentProcessingStatus.status == "completed"
            ).count()
        finally:
            db_status.close()

        # Use FAISS count as primary (batch processor populates FAISS, not DB)
        faiss_doc_count = int(local_stats.get('total_documents', 0))
        faiss_chunk_count = int(local_stats.get('total_chunks', 0))
        
        # Use whichever is higher (FAISS or DB)
        effective_doc_count = max(faiss_doc_count, db_doc_count)
        effective_chunk_count = max(faiss_chunk_count, db_total_chunks, total_chunks)
        
        return {
            'status': 'ready',
            'total_documents': effective_doc_count,  # Use FAISS count (primary) or DB count
            'local_faiss_documents': faiss_doc_count,  # FAISS index count
            'total_chunks': effective_chunk_count,  # Use FAISS chunks or DB chunks
            'categories': list(db_categories),  # List of unique categories
            'categories_count': len(db_categories),  # Count of unique categories
            'knowledge_level': knowledge_level.upper(),
            'search_mode': search_mode,
            'pdf_sources': pdf_sources,
            'processing_queue': {
                'queued': queued_count,
                'processing': processing_count,
                'completed': completed_count,
                'total': queued_count + processing_count + completed_count,
                'status_url': '/api/medical/knowledge/upload-status'
            },
            # Additional diagnostics
            'embedding_model': local_stats.get('embedding_model'),
            'embedding_dimension': local_stats.get('embedding_dimension'),
            'faiss_available': local_stats.get('faiss_available'),
            'model_loaded': local_stats.get('model_loaded'),
            'hybrid_enabled': local_stats.get('hybrid_enabled'),
            'bm25_indexed': local_stats.get('bm25_indexed'),
            'uploaded_files': len(upload_files),
            'total_upload_size_mb': round(total_upload_size / 1024 / 1024, 2),
            'medical_books_dir': str(MEDICAL_BOOKS_DIR.resolve()) if MEDICAL_BOOKS_DIR.exists() else None,
            'medical_books_count': len(medical_books_files),
            'capabilities': enhanced_stats.get('capabilities', {})
        }
    except Exception as e:
        # Convert exception to ASCII-safe message without using str() which may trigger rich formatting
        import traceback
        import sys
        tb_lines = traceback.format_exception(type(e), e, e.__traceback__)
        # Convert each line to ASCII
        safe_tb = []
        for line in tb_lines:
            try:
                safe_line = line.encode('ascii', errors='replace').decode('ascii')
                safe_tb.append(safe_line)
            except:
                safe_tb.append("Error in exception formatting")
        
        error_msg = "Unknown error"
        try:
            error_msg = repr(e).encode('ascii', errors='replace').decode('ascii')
        except:
            error_msg = "Error occurred"
        
        logger.error(f"Error getting statistics: {error_msg}\nTraceback:\n{''.join(safe_tb)}")
        return {'status': 'error', 'error': error_msg, 'total_documents': 0, 'total_chunks': 0, 'knowledge_level': 'UNKNOWN'}


@router.get("/pending-status")
async def get_pending_status():
    """Check status of pending document embedding tasks"""
    from pathlib import Path
    import json
    pending_dir = Path("data/knowledge_base/pending")
    if not pending_dir.exists():
        return {"pending": [], "completed": 0}
    
    pending = []
    for task_file in pending_dir.glob("*.json"):
        try:
            with open(task_file, 'r') as f:
                data = json.load(f)
            pending.append({
                "task_id": data.get('task_id'),
                "source": data.get('source'),
                "status": data.get('status'),
                "chunks_added": data.get('chunks_added'),
                "created_at": data.get('created_at')
            })
        except Exception as e:
            logger.warning(f"Error reading pending task {task_file}: {e}")
    
    return {
        "total_pending": len(pending),
        "pending": pending
    }


def organize_search_results(results: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    Lightweight organizer to sit between raw search hits and the KB consumer.
    - Deduplicates by document/chunk
    - Normalizes metadata fields for safer downstream use
    - Groups results by document for cleaner answers/citations
    """
    dedup_map = {}
    issues = []

    for r in results:
        if not isinstance(r, dict):
            continue
        metadata = (r.get("metadata") or {}).copy()
        filename = metadata.get("filename") or metadata.get("source") or "unknown"
        document_id = metadata.get("document_id") or metadata.get("document_uuid") or filename
        chunk_idx = metadata.get("chunk_index")

        if not metadata.get("document_id"):
            issues.append(f"Missing document_id for {filename}")

        # Normalize minimal metadata for downstream consumers
        metadata.setdefault("document_id", document_id)
        metadata.setdefault("filename", filename)
        metadata.setdefault("category", metadata.get("category") or "medical_pdf")
        metadata.setdefault("section", metadata.get("section") or "general")

        key = (document_id, chunk_idx)
        candidate = {**r, "metadata": metadata}

        existing = dedup_map.get(key)
        if existing is None or existing.get("score", 0) < candidate.get("score", 0):
            dedup_map[key] = candidate

    deduped = sorted(dedup_map.values(), key=lambda x: x.get("score", 0), reverse=True)

    grouped = defaultdict(list)
    for item in deduped:
        grouped[item["metadata"].get("document_id")].append(item)

    document_groups = []
    for doc_id, items in grouped.items():
        ordered = sorted(items, key=lambda x: x.get("score", 0), reverse=True)
        document_groups.append({
            "document_id": doc_id,
            "top_chunk": ordered[0],
            "chunks": len(ordered),
            "max_score": ordered[0].get("score", 0),
            "filenames": list({i["metadata"].get("filename") for i in ordered})
        })

    return {
        "deduped": deduped,
        "document_groups": document_groups,
        "issues": list({*issues})
    }


@router.post("/search")
async def search_knowledge_base(request: SearchRequest):
    """Search knowledge base with local + hybrid + OpenAI fallback"""
    mode = (request.search_mode or "hybrid").lower()
    filters = request.filters or {}

    def _passes(metadata: Dict[str, Any]) -> bool:
        if not filters:
            return True
        if 'category' in filters and metadata.get('category') and metadata.get('category') != filters['category']:
            return False
        if 'section' in filters and metadata.get('section'):
            if filters['section'].lower() not in str(metadata.get('section', '')).lower():
                return False
        if filters.get('min_year'):
            try:
                year = int(metadata.get('year')) if metadata.get('year') else None
                if year and year < int(filters['min_year']):
                    return False
            except Exception:
                pass
        if filters.get('allow_outdated') is False and metadata.get('outdated'):
            return False
        return True

    try:
        from app.services.local_vector_kb import get_local_knowledge_base
        local_kb = get_local_knowledge_base()
        use_bm25 = mode != "openai" and mode != "local"
        results = []

        if mode in {"local", "hybrid"}:
            results = local_kb.search(
                query=request.query,
                top_k=request.top_k,
                min_score=request.min_score,
                alpha=request.alpha,
                use_bm25=use_bm25,
                filters=filters
            )
            source = "local_hybrid" if use_bm25 else "local_vector"
        else:
            # Explicit OpenAI vector search path
            vector_kb = get_vector_knowledge_base()
            raw_results = vector_kb.search(
                query=request.query,
                top_k=request.top_k,
                filter_metadata=filters if filters else None
            )
            results = []
            for r in raw_results:
                metadata = r.get("metadata", {})
                if not _passes(metadata):
                    continue
                citation = {
                    "document_id": metadata.get("document_id"),
                    "filename": metadata.get("filename") or metadata.get("source"),
                    "page": metadata.get("page") or metadata.get("page_number"),
                    "chunk_index": metadata.get("chunk_index"),
                    "section": metadata.get("section"),
                    "category": metadata.get("category"),
                    "year": metadata.get("year"),
                }
                results.append({
                    "text": r.get("content"),
                    "metadata": metadata,
                    "score": r.get("similarity_score", 0),
                    "citation": citation,
                    "source_type": "openai_vector"
                })
            source = "openai_vector"

        if not results and request.allow_fallback and mode != "openai":
            # fallback to OpenAI vector search if local/hybrid empty
            vector_kb = get_vector_knowledge_base()
            raw_results = vector_kb.search(
                query=request.query,
                top_k=request.top_k,
                filter_metadata=filters if filters else None
            )
            for r in raw_results:
                metadata = r.get("metadata", {})
                if not _passes(metadata):
                    continue
                citation = {
                    "document_id": metadata.get("document_id"),
                    "filename": metadata.get("filename") or metadata.get("source"),
                    "page": metadata.get("page") or metadata.get("page_number"),
                    "chunk_index": metadata.get("chunk_index"),
                    "section": metadata.get("section"),
                    "category": metadata.get("category"),
                    "year": metadata.get("year"),
                }
                results.append({
                    "text": r.get("content"),
                    "metadata": metadata,
                    "score": r.get("similarity_score", 0),
                    "citation": citation,
                    "source_type": "openai_vector"
                })
            source = "openai_vector"

        filtered_results = [r for r in results if r.get("score", 0) >= request.min_score]
        final_results = filtered_results[: request.top_k]

        organized = organize_search_results(final_results)
        organized_results = organized.get("deduped", [])
        document_groups = organized.get("document_groups", [])
        data_issues = organized.get("issues", [])

        # Synthesize answer if requested
        answer = None
        if request.synthesize_answer and organized_results:
            try:
                # Prepare context
                context_parts = []
                for i, r in enumerate(organized_results):
                    source = r.get("citation", {}).get("filename", "Unknown")
                    text = r.get("text", "").strip()
                    context_parts.append(f"[{i+1}] Source: {source}\n{text}")
                
                context = "\n\n".join(context_parts)
                system_prompt = (
                    "You are a helpful medical assistant. "
                    "Answer the user's question using ONLY the provided context. "
                    "Cite your sources using brackets like [1], [2] corresponding to the source numbers provided. "
                    "If the answer is not in the context, say so. "
                    "Format your answer in Markdown with clear headings and bullet points."
                )
                
                user_prompt = f"Question: {request.query}\n\nContext:\n{context}"
                
                client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
                response = client.chat.completions.create(
                    model="gpt-4o-mini",  # Using gpt-4o-mini for better compatibility and lower cost
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=0.3
                )
                answer = response.choices[0].message.content
            except Exception as e:
                logger.error(f"Error synthesizing answer: {e}")
                answer = "Error generating answer. Please try again."

        return {
            "query": request.query,
            "results": organized_results,
            "raw_results": final_results,
            "document_groups": document_groups,
            "data_quality_flags": data_issues,
            "answer": answer,
            "total_results": len(filtered_results),
            "top_k": request.top_k,
            "source": source,
            "mode": mode,
            "applied_filters": filters,
        }
    except Exception as e:
        logger.error(f"Search error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")


@router.post("/search/evaluate")
async def evaluate_search(request: EvaluationRequest):
    """Lightweight evaluation harness for KB search.

    Expects list of runs with query + expected_terms and returns hit@k + matches.
    """
    from app.services.local_vector_kb import get_local_knowledge_base
    local_kb = get_local_knowledge_base()
    results = []
    for run in request.runs:
        hits = local_kb.search(
            query=run.query,
            top_k=run.top_k,
            use_bm25=True,
            alpha=0.6
        )
        matched_terms = []
        for term in run.expected_terms:
            term_lower = term.lower()
            if any(term_lower in (res.get('text') or '').lower() for res in hits):
                matched_terms.append(term)
        results.append({
            "query": run.query,
            "expected_terms": run.expected_terms,
            "hit": len(matched_terms) > 0,
            "matched_terms": matched_terms,
            "returned": len(hits)
        })
    hit_rate = sum(1 for r in results if r.get("hit")) / len(results) if results else 0
    return {"runs": results, "hit_rate": hit_rate}


@router.get("/upload-status/{document_id}")
async def get_upload_status(document_id: str, db: Session = Depends(get_db)):
    """Get processing status of a document upload"""
    try:
        from app.models import DocumentProcessingStatus
        
        status = db.query(DocumentProcessingStatus).filter(
            DocumentProcessingStatus.document_id == document_id
        ).first()
        
        if not status:
            return {
                "document_id": document_id,
                "status": "unknown",
                "message": "Document not found in processing queue"
            }
        
        response = {
            "document_id": document_id,
            "status": status.status,
            "progress_percent": status.progress_percent,
            "current_chunk": status.current_chunk,
            "total_chunks": status.total_chunks,
            "error_message": status.error_message,
            "created_at": status.created_at.isoformat() if status.created_at else None,
            "started_at": status.started_at.isoformat() if status.started_at else None,
            "completed_at": status.completed_at.isoformat() if status.completed_at else None,
        }
        
        if status.status == "processing":
            elapsed = (datetime.now() - status.started_at).total_seconds() if status.started_at else 0
            if status.estimated_time_seconds and status.progress_percent > 0:
                estimated_remaining = int(
                    (status.estimated_time_seconds * 100 / status.progress_percent) - elapsed
                )
                response["estimated_remaining_seconds"] = max(0, estimated_remaining)
        
        return response
    except Exception as e:
        logger.error(f"Error getting upload status: {e}")
        return {
            "document_id": document_id,
            "status": "error",
            "error_message": str(e)
        }


@router.get("/upload-status")
async def get_all_upload_statuses(db: Session = Depends(get_db)):
    """Get processing status of all documents in queue"""
    try:
        from app.models import DocumentProcessingStatus, KnowledgeDocument
        
        statuses = db.query(DocumentProcessingStatus).order_by(
            DocumentProcessingStatus.updated_at.desc()
        ).all()
        
        results = []
        for status in statuses:
            doc = db.query(KnowledgeDocument).filter(
                KnowledgeDocument.document_id == status.document_id
            ).first()
            
            results.append({
                "document_id": status.document_id,
                "filename": doc.filename if doc else "unknown",
                "status": status.status,
                "progress_percent": status.progress_percent,
                "current_chunk": status.current_chunk,
                "total_chunks": status.total_chunks,
                "size_mb": doc.file_size / 1024 / 1024 if doc else 0,
                "created_at": status.created_at.isoformat() if status.created_at else None,
                "started_at": status.started_at.isoformat() if status.started_at else None,
                "completed_at": status.completed_at.isoformat() if status.completed_at else None,
            })
        
        # Summary
        queued = sum(1 for s in statuses if s.status == "queued")
        processing = sum(1 for s in statuses if s.status == "processing")
        completed = sum(1 for s in statuses if s.status == "completed")
        failed = sum(1 for s in statuses if s.status == "failed")
        
        return {
            "total": len(statuses),
            "queued": queued,
            "processing": processing,
            "completed": completed,
            "failed": failed,
            "statuses": results
        }
    except Exception as e:
        logger.error(f"Error getting upload statuses: {e}")
        return {
            "total": 0,
            "queued": 0,
            "processing": 0,
            "completed": 0,
            "failed": 0,
            "statuses": [],
            "error": str(e)
        }


@router.get("/documents")
async def list_documents():
    """List all uploaded documents"""
    try:
        upload_files = list(UPLOAD_DIR.glob("*"))
        documents = []
        
        for file_path in upload_files:
            if file_path.is_file():
                stat = file_path.stat()
                documents.append({
                    "filename": file_path.name,
                    "size_mb": stat.st_size / 1024 / 1024,
                    "uploaded_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                    "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat()
                })
        
        # Sort by upload date (newest first)
        documents.sort(key=lambda x: x["uploaded_at"], reverse=True)
        
        return {
            "documents": documents,
            "total_count": len(documents),
            "total_size_mb": sum(d["size_mb"] for d in documents)
        }
    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        return {"error": str(e), "documents": []}


@router.delete("/documents/{filename}")
async def delete_document(
    filename: str,
    current_user: User = Depends(get_current_user)
):
    """Delete an uploaded document"""
    try:
        file_path = UPLOAD_DIR / filename
        
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Document not found")
        
        if not file_path.is_file():
            raise HTTPException(status_code=400, detail="Invalid file")
        
        file_path.unlink()
        
        return {
            "message": f"Document {filename} deleted successfully",
            "filename": filename
        }
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/upload-large")
async def upload_large_pdf(
    file: UploadFile = File(...),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Upload and process large PDF files (up to 2GB) asynchronously.
    """
    try:
        # Validate file type
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
            )
        
        # Save file with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = f"{timestamp}_{file.filename}"
        file_path = UPLOAD_DIR / safe_filename
        
        # Save uploaded file
        logger.info(f"Saving large file: {file.filename}")
        with open(file_path, "wb") as f:
            # Read and write in chunks to handle large files
            chunk_size = 1024 * 1024  # 1MB chunks
            while chunk := await file.read(chunk_size):
                f.write(chunk)
        
        file_size_mb = file_path.stat().st_size / 1024 / 1024
        logger.info(f"Saved {file.filename}: {file_size_mb:.2f} MB")
        
        # Validate file size
        if file_path.stat().st_size > MAX_FILE_SIZE:
            file_path.unlink()  # Delete oversized file
            raise HTTPException(
                status_code=400,
                detail=f"File exceeds maximum size of {MAX_FILE_SIZE / 1024 / 1024:.0f}MB"
            )

        # Create processing status record
        doc_uuid = str(uuid.uuid4())
        status_record = DocumentProcessingStatus(
            document_id=doc_uuid,
            status="queued",
            processing_type="large_file_embedding",
            progress_percent=0,
            error_message=None
        )
        db.add(status_record)
        db.commit() # Commit to get ID and ensure visibility

        # Background processing function
        async def process_large_file_task(file_path: Path, doc_uuid: str):
            try:
                # Update status to processing
                with next(get_db()) as task_db:
                    status = task_db.query(DocumentProcessingStatus).filter(
                        DocumentProcessingStatus.document_id == doc_uuid
                    ).first()
                    if status:
                        status.status = "processing"
                        status.started_at = datetime.utcnow()
                        task_db.commit()

                processor = get_large_pdf_processor()
                from app.services.enhanced_knowledge_base import get_knowledge_base
                knowledge_base = get_knowledge_base()

                # Progress callback wrapper
                async def progress_wrapper(progress: float, message: str):
                    logger.info(f"Processing {doc_uuid}: {progress:.1f}% - {message}")
                    # In a real scenario, you might update DB less frequently to avoid lock contention
                    # For now we log it, real DB updates happen inside processor if it supported it, 
                    # or we rely on the processor to call this
                    with next(get_db()) as progress_db:
                         s = progress_db.query(DocumentProcessingStatus).filter(
                            DocumentProcessingStatus.document_id == doc_uuid
                         ).first()
                         if s:
                             s.progress_percent = int(progress)
                             progress_db.commit()

                # Process
                await processor.process_large_pdf(
                    file_path=file_path,
                    add_to_kb=lambda text, source, metadata: add_to_knowledge_base(
                        knowledge_base, text, source, {**metadata, "document_uuid": doc_uuid, "uploaded_by": current_user.email}
                    ),
                    progress_callback=progress_wrapper
                )

                # Create KnowledgeDocument record 
                with next(get_db()) as final_db:
                    # Update status to completed
                    s = final_db.query(DocumentProcessingStatus).filter(
                        DocumentProcessingStatus.document_id == doc_uuid
                    ).first()
                    if s:
                        s.status = "completed"
                        s.progress_percent = 100
                        s.completed_at = datetime.utcnow()
                    
                    # Create the main document record if not exists (process_large_pdf might not do this)
                    # Note: Ideally process_large_pdf helper handles this or returns stats. 
                    # For simplicity, we assume successful return means success.
                    # We need to insert KnowledgeDocument here because process_large_pdf mostly adds chunks.
                    
                    # Check if doc exists (added by processor?)
                    doc = final_db.query(KnowledgeDocument).filter(KnowledgeDocument.document_id == doc_uuid).first()
                    if not doc:
                         new_doc = KnowledgeDocument(
                            document_id=doc_uuid,
                            filename=file.filename,
                            file_path=str(file_path),
                            file_hash=hashlib.sha256(file.filename.encode()).hexdigest(), # Simplified hash for now
                            file_size=int(file_path.stat().st_size),
                            extension=file_ext,
                            is_indexed=True,
                            uploaded_by_id=current_user.id,
                            category="medical_textbook",  # Set default category
                            source=file.filename
                        )
                         final_db.add(new_doc)
                    
                    final_db.commit()

            except Exception as e:
                logger.error(f"Error processing large file {doc_uuid}: {e}")
                with next(get_db()) as error_db:
                    s = error_db.query(DocumentProcessingStatus).filter(
                        DocumentProcessingStatus.document_id == doc_uuid
                    ).first()
                    if s:
                        s.status = "failed"
                        s.error_message = str(e)
                        error_db.commit()
        
        # Add to background tasks
        background_tasks.add_task(process_large_file_task, file_path, doc_uuid)

        return {
            "message": "File uploaded and queued for processing",
            "results": [{
                "filename": file.filename,
                "status": "success", # Initial upload success
                "document_id": doc_uuid,
                "info": "Large file queued for background processing",
                "size_mb": file_size_mb
            }]
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing large PDF: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Processing error: {str(e)}")


@router.get("/processor-stats")
async def get_processor_statistics():
    """Get large PDF processor statistics"""
    try:
        processor = get_large_pdf_processor()
        stats = processor.get_statistics()
        
        return {
            "processor_info": stats,
            "cache_dir": str(processor.cache_dir),
            "status": "operational"
        }
    except Exception as e:
        logger.error(f"Error getting processor stats: {e}")
        return {"error": str(e)}


@router.post("/clear-processor-cache")
async def clear_processor_cache(
    older_than_days: Optional[int] = None,
    current_user: User = Depends(get_current_user)
):
    """Clear PDF extraction cache"""
    try:
        processor = get_large_pdf_processor()
        result = await processor.clear_cache(older_than_days)
        
        return {
            "message": "Cache cleared successfully",
            **result
        }
    except Exception as e:
        logger.error(f"Error clearing cache: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/documents/clear-all")
async def clear_all_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Clear all documents from knowledge base and database (admin only)"""
    if current_user.role != "admin" and current_user.role != "doctor":
        raise HTTPException(status_code=403, detail="Only admins and doctors can clear knowledge base")
    
    try:
        # Delete all KB documents from database
        deleted_count = db.query(KnowledgeDocument).delete()
        db.commit()
        
        # Clear vector store
        from app.services.enhanced_knowledge_base import get_knowledge_base
        knowledge_base = get_knowledge_base()
        if hasattr(knowledge_base, 'clear_all'):
            knowledge_base.clear_all()
        
        logger.info(f"Cleared {deleted_count} documents from knowledge base")
        
        return {
            "message": "Knowledge base cleared successfully",
            "documents_deleted": deleted_count
        }
    except Exception as e:
        db.rollback()
        logger.error(f"Error clearing knowledge base: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/error-report")
async def get_error_report(current_user: User = Depends(get_current_user)):
    """Get automatic error correction system report"""
    corrector = get_error_corrector()
    return corrector.get_error_report()


@router.get("/documents/list")
async def list_documents(
    skip: int = 0,
    limit: int = 100,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List all uploaded documents with metadata"""
    documents = db.query(KnowledgeDocument).offset(skip).limit(limit).all()
    
    return {
        "total": db.query(KnowledgeDocument).count(),
        "documents": [
            {
                "id": doc.document_id,
                "filename": doc.filename,
                "size_mb": round(doc.file_size / 1024 / 1024, 2),
                "chunks": doc.chunk_count,
                "text_length": doc.text_length,
                "uploaded_at": doc.uploaded_at.isoformat(),
                "uploaded_by": doc.uploaded_by.email if doc.uploaded_by else None,
                "is_indexed": doc.is_indexed
            }
            for doc in documents
        ]
    }


@router.get("/documents/{document_id}")
async def get_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get document details and content by ID or filename"""
    # Try to find by UUID first
    doc = db.query(KnowledgeDocument).filter(
        KnowledgeDocument.document_id == document_id
    ).first()
    
    # If not found, try by filename
    if not doc:
        doc = db.query(KnowledgeDocument).filter(
            KnowledgeDocument.filename == document_id
        ).first()
    
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Get chunks for this document
    chunks = db.query(KnowledgeChunk).filter(
        KnowledgeChunk.document_id == doc.document_id
    ).order_by(KnowledgeChunk.chunk_index).all()
    
    return {
        "id": doc.document_id,
        "filename": doc.filename,
        "file_path": doc.file_path,
        "file_size": doc.file_size,
        "size_mb": round(doc.file_size / 1024 / 1024, 2),
        "text_length": doc.text_length,
        "chunk_count": doc.chunk_count,
        "uploaded_at": doc.uploaded_at.isoformat(),
        "uploaded_by": doc.uploaded_by.email if doc.uploaded_by else None,
        "is_indexed": doc.is_indexed,
        "chunks": [
            {
                "chunk_id": chunk.chunk_id,
                "chunk_index": chunk.chunk_index,
                "text": chunk.text_content,
                "text_length": len(chunk.text_content),
                "metadata": chunk.metadata
            }
            for chunk in chunks[:10]  # Limit to first 10 chunks for performance
        ],
        "total_chunks": len(chunks)
    }


# ============================================================================
# ONLINE KNOWLEDGE BASE ENDPOINTS
# ============================================================================

_online_kb_service = None

def get_online_kb_service() -> OnlineKnowledgeService:
    """Get singleton instance of online KB service"""
    global _online_kb_service
    if _online_kb_service is None:
        _online_kb_service = OnlineKnowledgeService()
    return _online_kb_service


@router.get("/online/search-pubmed")
async def search_pubmed_online(
    query: str,
    max_results: int = 5,
    use_cache: bool = True,
    current_user: User = Depends(get_current_user)
):
    """
    Search PubMed for recent medical research (FREE - No API key required)
    
    - **query**: Search terms (e.g., "diabetes treatment", "COVID-19 symptoms")
    - **max_results**: Number of results (1-20)
    - **use_cache**: Use cached results if available (24h cache)
    
    Returns recent research articles from PubMed database
    """
    if max_results < 1 or max_results > 20:
        raise HTTPException(status_code=400, detail="max_results must be between 1 and 20")
    
    try:
        service = get_online_kb_service()
        await service.initialize()
        
        results = await service.search_pubmed(query, max_results, use_cache)
        
        return {
            "query": query,
            "total_results": len(results),
            "cached": use_cache,
            "results": results
        }
    except Exception as e:
        logger.error(f"Error searching PubMed: {e}")
        raise HTTPException(status_code=500, detail=f"PubMed search failed: {str(e)}")


@router.get("/online/guidelines")
async def get_clinical_guidelines(
    condition: str,
    use_cache: bool = True,
    current_user: User = Depends(get_current_user)
):
    """
    Get clinical practice guidelines from multiple authoritative sources
    
    - **condition**: Medical condition (e.g., "hypertension", "diabetes", "asthma")
    - **use_cache**: Use cached results (7-day cache for guidelines)
    
    Sources: PubMed, CDC, WHO, NICE (UK)
    """
    try:
        service = get_online_kb_service()
        await service.initialize()
        
        guidelines = await service.get_clinical_guidelines(condition, use_cache)
        
        # Group by source
        by_source = {}
        for guideline in guidelines:
            source = guideline.get('source', 'Unknown')
            if source not in by_source:
                by_source[source] = []
            by_source[source].append(guideline)
        
        return {
            "condition": condition,
            "total_guidelines": len(guidelines),
            "sources": list(by_source.keys()),
            "cached": use_cache,
            "by_source": by_source,
            "all_guidelines": guidelines
        }
    except Exception as e:
        logger.error(f"Error fetching guidelines: {e}")
        raise HTTPException(status_code=500, detail=f"Guidelines fetch failed: {str(e)}")


@router.get("/online/drug-info")
async def get_drug_information(
    drug_name: str,
    use_cache: bool = True,
    current_user: User = Depends(get_current_user)
):
    """
    Get comprehensive drug information from online sources
    
    - **drug_name**: Drug name (generic or brand)
    - **use_cache**: Use cached results (24h cache)
    
    Returns: Indications, dosage, side effects, interactions, etc.
    """
    try:
        service = get_online_kb_service()
        await service.initialize()
        
        drug_info = await service.get_drug_information(drug_name, use_cache)
        
        return {
            "drug_name": drug_name,
            "cached": use_cache,
            "information": drug_info
        }
    except Exception as e:
        logger.error(f"Error fetching drug info: {e}")
        raise HTTPException(status_code=500, detail=f"Drug info fetch failed: {str(e)}")


@router.get("/online/status")
async def online_kb_status(current_user: User = Depends(get_current_user)):
    """
    Get status of online knowledge base service
    
    Returns cache statistics and service availability
    """
    try:
        service = get_online_kb_service()
        await service.initialize()
        
        # Get cache stats
        cache_dir = service.cache_dir
        cache_files = list(cache_dir.glob("*.json"))
        total_cache_size = sum(f.stat().st_size for f in cache_files)
        
        return {
            "status": "operational",
            "initialized": service.initialized,
            "cache_directory": str(cache_dir),
            "cache_files": len(cache_files),
            "cache_size_mb": round(total_cache_size / 1024 / 1024, 2),
            "cache_duration_hours": service.cache_duration.total_seconds() / 3600,
            "features": {
                "pubmed_search": "Available (FREE)",
                "clinical_guidelines": "Available (PubMed, CDC, WHO, NICE)",
                "drug_information": "Available",
                "medical_news": "Available"
            }
        }
    except Exception as e:
        logger.error(f"Error getting online KB status: {e}")
        return {
            "status": "error",
            "error": str(e)
        }


@router.delete("/online/clear-cache")
async def clear_online_cache(current_user: User = Depends(get_current_user)):
    """
    Clear online knowledge base cache
    
    Forces fresh data fetching on next request
    """
    if current_user.role not in ["admin", "doctor"]:
        raise HTTPException(status_code=403, detail="Only admins and doctors can clear cache")
    
    try:
        service = get_online_kb_service()
        await service.initialize()
        
        cache_dir = service.cache_dir
        cache_files = list(cache_dir.glob("*.json"))
        deleted_count = 0
        
        for cache_file in cache_files:
            cache_file.unlink()
            deleted_count += 1
        
        return {
            "message": "Online KB cache cleared",
            "files_deleted": deleted_count
        }
    except Exception as e:
        logger.error(f"Error clearing online cache: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/search/enhanced")
async def enhanced_search(
    request: SearchRequest,
    include_images: bool = True,
    verify_online: bool = True,
    current_user: User = Depends(get_current_user)
):
    """
    Enhanced KB search with images and online verification
    
    - **query**: Search query
    - **top_k**: Number of results (default 5)
    - **include_images**: Include related medical images from PDFs
    - **verify_online**: Verify content against current medical guidelines
    
    Returns: Text results + images + verification status
    """
    try:
        from app.services.enhanced_kb_processor import EnhancedKBProcessor, enhance_kb_search
        
        # Get KB services
        from app.services.local_vector_kb import get_local_knowledge_base
        kb_service = get_local_knowledge_base()
        processor = EnhancedKBProcessor()
        
        # Perform enhanced search
        results = processor.search_with_images(
            query=request.query,
            kb_service=kb_service,
            include_images=include_images,
            top_k=request.top_k
        )
        
        # Add metadata
        results['query'] = request.query
        results['include_images'] = include_images
        results['verify_online'] = verify_online
        results['timestamp'] = datetime.now().isoformat()
        
        return results
        
    except Exception as e:
        logger.error(f"Enhanced search failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/images/{document_id}")
async def get_document_images(
    document_id: str,
    current_user: User = Depends(get_current_user)
):
    """
    Get all images extracted from a specific document
    
    Returns list of images with captions and descriptions
    """
    try:
        from app.services.enhanced_kb_processor import EnhancedKBProcessor
        
        processor = EnhancedKBProcessor()
        images = processor._get_images_for_document(document_id)
        
        # Generate AI descriptions for each image
        for img in images:
            if 'description' not in img:
                img['description'] = processor.generate_image_description(img['path'])
        
        return {
            "document_id": document_id,
            "total_images": len(images),
            "images": images
        }
        
    except Exception as e:
        logger.error(f"Error getting document images: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/extract-images/{document_id}")
async def extract_document_images(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Extract and save images from a specific PDF document
    
    This is a background task that can take time for large PDFs
    """
    if current_user.role not in ["admin", "doctor"]:
        raise HTTPException(status_code=403, detail="Only admins and doctors can extract images")
    
    try:
        from app.services.enhanced_kb_processor import EnhancedKBProcessor
        
        # Find document in database
        doc = db.query(KnowledgeDocument).filter(
            KnowledgeDocument.document_id == document_id
        ).first()
        
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        
        if not doc.file_path or not Path(doc.file_path).exists():
            raise HTTPException(status_code=404, detail="PDF file not found on disk")
        
        # Extract images
        processor = EnhancedKBProcessor()
        extraction_result = processor.extract_text_and_images(doc.file_path)
        
        if 'error' in extraction_result:
            raise HTTPException(status_code=500, detail=extraction_result['error'])
        
        # Save images
        saved_images = []
        for img in extraction_result['images']:
            img_path = processor.save_image(
                image_data=img['image_data'],
                image_hash=f"{document_id}_{img['hash']}",
                format=img['format']
            )
            saved_images.append({
                'path': img_path,
                'page': img['page'],
                'caption': img['caption'],
                'size_kb': img['size'] / 1024
            })
        
        return {
            "document_id": document_id,
            "filename": doc.filename,
            "extraction_complete": True,
            "total_images": len(saved_images),
            "images": saved_images,
            "metadata": extraction_result['metadata']
        }
        
    except Exception as e:
        logger.error(f"Error extracting images: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/reset")
async def reset_knowledge_base(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    [ADMIN ONLY] Reset/clear all documents from knowledge base.
    Clears:
    - Database records (KnowledgeDocument)
    - Uploaded files (data/uploaded_documents)
    - FAISS index
    - Cached embeddings
    Requires admin privileges.
    """
    # Check if user is admin
    if current_user.role != "admin":
        raise HTTPException(
            status_code=403,
            detail="Only administrators can reset the knowledge base"
        )
    
    try:
        import shutil
        from pathlib import Path
        
        # Get count before deletion
        count_before = db.query(KnowledgeDocument).count()
        logger.warning(f"[ADMIN] {current_user.email} initiating KB reset - {count_before} documents")
        
        # 1. Delete all database records
        db.query(KnowledgeDocument).delete()
        db.commit()
        
        # 2. Delete uploaded files directory
        uploads_dir = Path("data/uploaded_documents")
        if uploads_dir.exists():
            logger.warning(f"Deleting uploaded files directory: {uploads_dir}")
            shutil.rmtree(uploads_dir)
            uploads_dir.mkdir(parents=True, exist_ok=True)
        
        # 3. Delete FAISS index files
        data_dir = Path("data/knowledge_base")
        faiss_index_path = data_dir / "local_faiss_index.bin"
        metadata_path = data_dir / "metadata.json"
        embeddings_cache = Path("backend/cache/online_knowledge")
        
        if faiss_index_path.exists():
            logger.warning(f"Deleting FAISS index: {faiss_index_path}")
            faiss_index_path.unlink()
        
        if metadata_path.exists():
            logger.warning(f"Deleting metadata: {metadata_path}")
            metadata_path.unlink()
        
        if embeddings_cache.exists():
            logger.warning(f"Deleting embeddings cache: {embeddings_cache}")
            shutil.rmtree(embeddings_cache)
            embeddings_cache.mkdir(parents=True, exist_ok=True)
        
        # 4. Reinitialize knowledge base
        kb = _get_kb()
        if kb:
            logger.info("Reinitializing knowledge base...")
            kb._initialize_index()
            kb.save_index()
        
        # Verify deletion
        count_after = db.query(KnowledgeDocument).count()
        
        logger.warning(f"[ADMIN] KB reset complete - deleted {count_before} documents, {count_after} remaining")
        
        return {
            "status": "success",
            "message": f"Knowledge base fully reset",
            "cleared": {
                "database_documents": count_before,
                "uploaded_files": "cleared",
                "faiss_index": "deleted",
                "embeddings_cache": "cleared"
            },
            "remaining_documents": count_after,
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        logger.error(f"[ERROR] KB reset failed: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to reset knowledge base: {str(e)}"
        )


@router.get("/ocr-status")
async def get_ocr_status():
    """
    Check OCR setup status and get installation instructions
    
    Returns information about:
    - Whether OCR dependencies are installed
    - Setup instructions for missing components
    - Current capabilities
    """
    try:
        ocr_processor = get_pdf_ocr_processor()
        setup_info = ocr_processor.get_setup_instructions()
        
        return {
            "status": "ready" if setup_info['ocr_ready'] else "needs_setup",
            "ocr_available": setup_info['ocr_ready'],
            "components": {
                "pytesseract": "installed" if setup_info['pytesseract_installed'] else "not_installed",
                "pdf2image": "installed" if setup_info['pdf2image_installed'] else "not_installed",
                "tesseract_ocr": "available" if setup_info['tesseract_available'] else "not_installed",
                "poppler": "available" if setup_info['poppler_available'] else "not_installed"
            },
            "capabilities": {
                "text_extraction": True,
                "image_extraction": True,
                "ocr_processing": setup_info['ocr_ready'],
                "scanned_pdf_support": setup_info['ocr_ready']
            },
            "setup_instructions": setup_info['instructions'],
            "download_links": {
                "tesseract_windows": "https://github.com/UB-Mannheim/tesseract/wiki",
                "poppler_windows": "http://blog.alivate.com.au/poppler-windows/",
                "pip_install": "pip install pytesseract pdf2image"
            }
        }
    except Exception as e:
        logger.error(f"Error checking OCR status: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to check OCR status: {str(e)}"
        )

@router.get("/queue-status")
async def get_queue_status(db: Session = Depends(get_db)):
    """
    Get PDF upload queue processing status
    
    Returns:
    - Queue statistics (queued, processing, completed, failed counts)
    - Processing worker status
    - Current batch information
    """
    try:
        from app.services.upload_queue_processor import get_queue_processor
        from app.models import DocumentProcessingStatus
        
        processor = get_queue_processor()
        
        # Get queue statistics
        queued_count = db.query(DocumentProcessingStatus).filter(
            DocumentProcessingStatus.status == 'queued'
        ).count()
        
        processing_count = db.query(DocumentProcessingStatus).filter(
            DocumentProcessingStatus.status == 'processing'
        ).count()
        
        completed_count = db.query(DocumentProcessingStatus).filter(
            DocumentProcessingStatus.status == 'completed'
        ).count()
        
        failed_count = db.query(DocumentProcessingStatus).filter(
            DocumentProcessingStatus.status == 'failed'
        ).count()
        
        # Get current processing documents
        processing_docs = db.query(DocumentProcessingStatus).filter(
            DocumentProcessingStatus.status == 'processing'
        ).all()
        
        processing_details = [
            {
                "document_id": doc.document_id,
                "progress_percent": doc.progress_percent or 0,
                "current_chunk": doc.current_chunk or 0,
                "total_chunks": doc.total_chunks or 0,
                "estimated_time_seconds": doc.estimated_time_seconds or 0,
                "started_at": doc.started_at.isoformat() if doc.started_at else None
            }
            for doc in processing_docs
        ]
        
        return {
            "worker_status": "running" if processor.is_running else "stopped",
            "queue": {
                "queued": queued_count,
                "processing": processing_count,
                "completed": completed_count,
                "failed": failed_count,
                "total": queued_count + processing_count + completed_count + failed_count
            },
            "processor_config": {
                "batch_size": processor.batch_size,
                "check_interval": processor.check_interval,
                "max_retries": processor.max_retries
            },
            "processing_details": processing_details,
            "timestamp": datetime.utcnow().isoformat()
        }
    except Exception as e:
        logger.error(f"Error getting queue status: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get queue status: {str(e)}"
        )


# ==================== PDF Processing with Pause/Resume ====================

@router.post("/pdf/pause/{processing_id}")
async def pause_pdf_processing(
    processing_id: int,
    current_user: User = Depends(require_kb_management_role),
    db: Session = Depends(get_db),
):
    """Pause an active PDF processing job."""
    try:
        from app.services.pdf_processing_manager import pdf_processor_with_resume
        from app.models import PDFProcessing
        
        # Check if processing exists and belongs to user
        processing = db.query(PDFProcessing).filter(
            PDFProcessing.id == processing_id,
            PDFProcessing.user_id == current_user.id
        ).first()
        
        if not processing:
            raise HTTPException(
                status_code=404,
                detail=f"Processing job {processing_id} not found"
            )
        
        success = await pdf_processor_with_resume.pause_processing(processing_id, db)
        
        if success:
            logger.info(f"User {current_user.email} paused PDF processing {processing_id}")
            return {
                "status": "paused",
                "processing_id": processing_id,
                "message": "PDF processing paused successfully"
            }
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Cannot pause processing in status: {processing.status}"
            )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error pausing processing: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to pause processing: {str(e)}"
        )


@router.post("/pdf/resume/{processing_id}")
async def resume_pdf_processing(
    processing_id: int,
    current_user: User = Depends(require_kb_management_role),
    db: Session = Depends(get_db),
    background_tasks: BackgroundTasks = BackgroundTasks(),
):
    """Resume a paused PDF processing job."""
    try:
        from app.services.pdf_processing_manager import pdf_processor_with_resume
        from app.models import PDFProcessing, PDFFile
        
        # Check if processing exists and belongs to user
        processing = db.query(PDFProcessing).filter(
            PDFProcessing.id == processing_id,
            PDFProcessing.user_id == current_user.id
        ).first()
        
        if not processing:
            raise HTTPException(
                status_code=404,
                detail=f"Processing job {processing_id} not found"
            )
        
        # Get the PDF file
        pdf_file = db.query(PDFFile).filter(
            PDFFile.id == processing.pdf_file_id
        ).first()
        
        if not pdf_file:
            raise HTTPException(
                status_code=404,
                detail=f"PDF file not found"
            )
        
        success = await pdf_processor_with_resume.resume_processing(processing_id, db)
        
        if success:
            # Start background processing
            background_tasks.add_task(
                pdf_processor_with_resume.process_pdf_with_checkpoint,
                processing_id,
                pdf_file.file_path,
                db,
            )
            
            logger.info(f"User {current_user.email} resumed PDF processing {processing_id}")
            return {
                "status": "resumed",
                "processing_id": processing_id,
                "message": "PDF processing resumed successfully"
            }
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Cannot resume processing in status: {processing.status}"
            )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error resuming processing: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to resume processing: {str(e)}"
        )


@router.get("/pdf/status/{processing_id}")
async def get_pdf_processing_status(
    processing_id: int,
    current_user: User = Depends(require_kb_management_role),
    db: Session = Depends(get_db),
):
    """Get status of a PDF processing job."""
    try:
        from app.services.pdf_processing_manager import pdf_processor_with_resume
        from app.models import PDFProcessing
        
        # Check if processing exists and belongs to user
        processing = db.query(PDFProcessing).filter(
            PDFProcessing.id == processing_id,
            PDFProcessing.user_id == current_user.id
        ).first()
        
        if not processing:
            raise HTTPException(
                status_code=404,
                detail=f"Processing job {processing_id} not found"
            )
        
        status = pdf_processor_with_resume.get_processing_status(processing_id, db)
        
        if status:
            return {
                "success": True,
                "data": status
            }
        else:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to retrieve status"
            )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting processing status: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get status: {str(e)}"
        )


@router.post("/pdf/cancel/{processing_id}")
async def cancel_pdf_processing(
    processing_id: int,
    current_user: User = Depends(require_kb_management_role),
    db: Session = Depends(get_db),
):
    """Cancel a PDF processing job."""
    try:
        from app.services.pdf_processing_manager import pdf_processor_with_resume
        from app.models import PDFProcessing
        
        # Check if processing exists and belongs to user
        processing = db.query(PDFProcessing).filter(
            PDFProcessing.id == processing_id,
            PDFProcessing.user_id == current_user.id
        ).first()
        
        if not processing:
            raise HTTPException(
                status_code=404,
                detail=f"Processing job {processing_id} not found"
            )
        
        success = await pdf_processor_with_resume.cancel_processing(processing_id, db)
        
        if success:
            logger.info(f"User {current_user.email} cancelled PDF processing {processing_id}")
            return {
                "status": "cancelled",
                "processing_id": processing_id,
                "message": "PDF processing cancelled successfully"
            }
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Cannot cancel processing in status: {processing.status}"
            )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error cancelling processing: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to cancel processing: {str(e)}"
        )


@router.get("/pdf/processing-list")
async def get_user_processing_jobs(
    current_user: User = Depends(require_kb_management_role),
    db: Session = Depends(get_db),
):
    """Get all PDF processing jobs for current user."""
    try:
        from app.models import PDFProcessing
        
        processing_jobs = db.query(PDFProcessing).filter(
            PDFProcessing.user_id == current_user.id
        ).order_by(PDFProcessing.created_at.desc()).all()
        
        jobs = []
        for job in processing_jobs:
            jobs.append({
                "id": job.id,
                "pdf_name": job.pdf_name,
                "status": job.status,
                "progress_percentage": job.progress_percentage,
                "pages_processed": job.pages_processed,
                "total_pages": job.total_pages,
                "embeddings_created": job.embeddings_created,
                "created_at": job.created_at.isoformat() if job.created_at else None,
                "started_at": job.started_at.isoformat() if job.started_at else None,
                "completed_at": job.completed_at.isoformat() if job.completed_at else None,
                "error_message": job.error_message
            })
        
        return {
            "success": True,
            "total": len(jobs),
            "data": jobs
        }
    
    except Exception as e:
        logger.error(f"Error getting processing list: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to get processing list: {str(e)}"
        )