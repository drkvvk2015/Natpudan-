"""
Enhanced Document Manager with Optimized PDF Processing
Implements async processing, parallel extraction, caching, and error recovery
"""

import logging
import os
import hashlib
import asyncio
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from datetime import datetime
import uuid
import json
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
import time

logger = logging.getLogger(__name__)

# Optional imports with fallbacks
try:
    import PyPDF2
    from PyPDF2.errors import PdfReadError
    PDF_AVAILABLE = True
except ImportError:
    logger.warning("PyPDF2 not available. Install with: pip install PyPDF2")
    PDF_AVAILABLE = False
    PdfReadError = Exception

try:
    import fitz  # PyMuPDF - faster alternative
    PYMUPDF_AVAILABLE = True
except ImportError:
    logger.warning("PyMuPDF not available. Install with: pip install PyMuPDF")
    PYMUPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available. Install with: pip install python-docx")
    DOCX_AVAILABLE = False


class EnhancedDocumentManager:
    """
    Enhanced document manager with:
    - Async/parallel processing
    - Multiple PDF extraction engines (PyMuPDF, PyPDF2)
    - Intelligent caching
    - Comprehensive error handling
    - Progress tracking
    - Performance monitoring
    """
    
    def __init__(
        self,
        upload_dir: str = "data/uploaded_documents",
        cache_dir: str = "data/document_cache",
        allowed_extensions: List[str] = None,
        max_workers: int = 4,
        enable_cache: bool = True
    ):
        """
        Initialize enhanced document manager.
        
        Args:
            upload_dir: Directory to store uploaded files
            cache_dir: Directory for caching extracted text
            allowed_extensions: Allowed file extensions
            max_workers: Maximum parallel workers
            enable_cache: Enable text extraction caching
        """
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        
        if allowed_extensions is None:
            self.allowed_extensions = ['.pdf', '.docx', '.txt', '.md']
        else:
            self.allowed_extensions = allowed_extensions
        
        self.max_workers = max_workers
        self.enable_cache = enable_cache
        
        # Performance metrics
        self.metrics = {
            "total_processed": 0,
            "cache_hits": 0,
            "cache_misses": 0,
            "errors": 0,
            "total_processing_time": 0.0
        }
        
        logger.info(f"Enhanced document manager initialized. Upload dir: {self.upload_dir}")
    
    def is_allowed_file(self, filename: str) -> bool:
        """Check if file extension is allowed"""
        return any(filename.lower().endswith(ext) for ext in self.allowed_extensions)
    
    async def save_upload(
        self,
        file_content: bytes,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None,
        max_file_size_mb: int = 500
    ) -> Dict[str, Any]:
        """
        Save uploaded file with async processing and error handling.
        
        Args:
            file_content: File content bytes
            filename: Original filename
            metadata: Additional metadata
            max_file_size_mb: Maximum file size in MB (default: 500MB)
            
        Returns:
            Document information dictionary
        """
        start_time = time.time()
        
        try:
            # Validate file type
            if not self.is_allowed_file(filename):
                raise ValueError(f"File type not allowed: {filename}. Allowed: {self.allowed_extensions}")
            
            # Validate file size with configurable limit
            max_bytes = max_file_size_mb * 1024 * 1024
            if len(file_content) > max_bytes:
                raise ValueError(
                    f"File too large: {len(file_content) / (1024 * 1024):.1f}MB. "
                    f"Maximum: {max_file_size_mb}MB"
                )
            
            # Generate unique document ID
            document_id = str(uuid.uuid4())
            file_hash = hashlib.sha256(file_content).hexdigest()
            
            # Check for duplicate uploads
            duplicate = await self._check_duplicate(file_hash)
            if duplicate:
                logger.info(f"Duplicate file detected: {filename} (hash: {file_hash[:8]}...)")
                return duplicate
            
            # Create document directory
            doc_dir = self.upload_dir / document_id
            doc_dir.mkdir(parents=True, exist_ok=True)
            
            # Save file
            file_path = doc_dir / filename
            async with asyncio.Lock():
                with open(file_path, 'wb') as f:
                    f.write(file_content)
            
            # Extract text asynchronously with progress tracking
            try:
                text_content = await self._extract_text_async(
                    file_path,
                    file_hash
                )
            except Exception as e:
                logger.error(f"Error extracting text from {filename}: {e}")
                text_content = ""
                self.metrics["errors"] += 1
            
            # Create document info
            processing_time = time.time() - start_time
            doc_info = {
                'document_id': document_id,
                'filename': filename,
                'file_path': str(file_path),
                'file_size': len(file_content),
                'file_hash': file_hash,
                'extension': Path(filename).suffix.lower(),
                'uploaded_at': datetime.utcnow().isoformat(),
                'text_length': len(text_content),
                'processing_time_seconds': round(processing_time, 2),
                'metadata': metadata or {},
                'status': 'success' if text_content else 'partial_failure'
            }
            
            # Save metadata
            metadata_path = doc_dir / 'metadata.json'
            with open(metadata_path, 'w', encoding='utf-8') as f:
                json.dump(doc_info, f, indent=2)
            
            # Save extracted text
            if text_content:
                text_path = doc_dir / 'extracted_text.txt'
                with open(text_path, 'w', encoding='utf-8') as f:
                    f.write(text_content)
            
            # Update metrics
            self.metrics["total_processed"] += 1
            self.metrics["total_processing_time"] += processing_time
            
            logger.info(
                f"Processed document: {document_id} ({filename}) "
                f"in {processing_time:.2f}s - {len(text_content)} chars extracted"
            )
            
            return doc_info
            
        except Exception as e:
            logger.error(f"Error saving upload {filename}: {e}")
            self.metrics["errors"] += 1
            raise
    
    async def _check_duplicate(self, file_hash: str) -> Optional[Dict[str, Any]]:
        """Check if file with same hash already exists"""
        for doc_dir in self.upload_dir.iterdir():
            if not doc_dir.is_dir():
                continue
            
            metadata_path = doc_dir / 'metadata.json'
            if metadata_path.exists():
                try:
                    with open(metadata_path, 'r', encoding='utf-8') as f:
                        doc_info = json.load(f)
                    
                    if doc_info.get('file_hash') == file_hash:
                        return doc_info
                except Exception as e:
                    logger.warning(f"Error reading metadata in {doc_dir}: {e}")
        
        return None
    
    async def _extract_text_async(
        self,
        file_path: Path,
        file_hash: str
    ) -> str:
        """
        Extract text asynchronously with caching.
        
        Args:
            file_path: Path to file
            file_hash: File hash for cache key
            
        Returns:
            Extracted text
        """
        # Check cache first
        if self.enable_cache:
            cached_text = await self._get_cached_text(file_hash)
            if cached_text is not None:
                self.metrics["cache_hits"] += 1
                logger.debug(f"Cache hit for {file_path.name}")
                return cached_text
            
            self.metrics["cache_misses"] += 1
        
        # Extract text
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                text = await self._extract_pdf_async(file_path)
            elif extension == '.docx':
                text = await self._extract_docx_async(file_path)
            elif extension in ['.txt', '.md']:
                text = await self._extract_text_file_async(file_path)
            else:
                raise ValueError(f"Unsupported file type: {extension}")
            
            # Cache the result
            if self.enable_cache and text:
                await self._cache_text(file_hash, text)
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    async def _get_cached_text(self, file_hash: str) -> Optional[str]:
        """Get cached extracted text"""
        cache_path = self.cache_dir / f"{file_hash}.txt"
        
        if cache_path.exists():
            try:
                loop = asyncio.get_event_loop()
                return await loop.run_in_executor(
                    None,
                    lambda: cache_path.read_text(encoding='utf-8')
                )
            except Exception as e:
                logger.warning(f"Error reading cache: {e}")
        
        return None
    
    async def _cache_text(self, file_hash: str, text: str) -> None:
        """Cache extracted text"""
        cache_path = self.cache_dir / f"{file_hash}.txt"
        
        try:
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(
                None,
                lambda: cache_path.write_text(text, encoding='utf-8')
            )
        except Exception as e:
            logger.warning(f"Error writing cache: {e}")
    
    async def _extract_pdf_async(self, file_path: Path) -> str:
        """
        Extract text from PDF with multiple engines and parallel processing.
        Tries PyMuPDF first (faster), falls back to PyPDF2.
        """
        # Try PyMuPDF first (much faster)
        if PYMUPDF_AVAILABLE:
            try:
                loop = asyncio.get_event_loop()
                text = await loop.run_in_executor(
                    None,
                    self._extract_pdf_pymupdf,
                    file_path
                )
                if text and len(text.strip()) > 50:  # Validate extraction
                    return text
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed, trying PyPDF2: {e}")
        
        # Fallback to PyPDF2
        if PDF_AVAILABLE:
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(
                None,
                self._extract_pdf_pypdf2,
                file_path
            )
        
        raise RuntimeError("No PDF extraction library available")
    
    def _extract_pdf_pymupdf(self, file_path: Path) -> str:
        """Extract PDF using PyMuPDF (faster)"""
        text_parts = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                try:
                    page = doc[page_num]
                    page_text = page.get_text()
                    
                    if page_text and page_text.strip():
                        text_parts.append(f"\n--- Page {page_num + 1} ---\n")
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num} with PyMuPDF: {e}")
            
            doc.close()
            
        except Exception as e:
            logger.error(f"PyMuPDF error: {e}")
            raise
        
        return '\n'.join(text_parts)
    
    def _extract_pdf_pypdf2(self, file_path: Path) -> str:
        """Extract PDF using PyPDF2 (fallback)"""
        text_parts = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        
                        if page_text and page_text.strip():
                            text_parts.append(f"\n--- Page {page_num + 1} ---\n")
                            text_parts.append(page_text)
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num} with PyPDF2: {e}")
        
        except PdfReadError as e:
            logger.error(f"PDF read error: {e}")
            raise ValueError(f"Corrupted or encrypted PDF: {e}")
        except Exception as e:
            logger.error(f"PyPDF2 error: {e}")
            raise
        
        return '\n'.join(text_parts)
    
    async def _extract_docx_async(self, file_path: Path) -> str:
        """Extract text from DOCX asynchronously"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed")
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            self._extract_docx,
            file_path
        )
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise
    
    async def _extract_text_file_async(self, file_path: Path) -> str:
        """Extract text from plain text file asynchronously"""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None,
            self._extract_text_file,
            file_path
        )
    
    def _extract_text_file(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with latin-1 encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    def get_document(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Get document information"""
        metadata_path = self.upload_dir / document_id / 'metadata.json'
        
        if metadata_path.exists():
            try:
                with open(metadata_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logger.error(f"Error reading metadata for {document_id}: {e}")
        
        return None
    
    def get_document_text(self, document_id: str) -> Optional[str]:
        """Get extracted text for a document"""
        text_path = self.upload_dir / document_id / 'extracted_text.txt'
        
        if text_path.exists():
            try:
                with open(text_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except Exception as e:
                logger.error(f"Error reading text for {document_id}: {e}")
        
        return None
    
    def list_documents(self) -> List[Dict[str, Any]]:
        """List all uploaded documents"""
        documents = []
        
        for doc_dir in self.upload_dir.iterdir():
            if doc_dir.is_dir():
                doc_info = self.get_document(doc_dir.name)
                if doc_info:
                    documents.append(doc_info)
        
        # Sort by upload date (newest first)
        documents.sort(
            key=lambda x: x.get('uploaded_at', ''),
            reverse=True
        )
        
        return documents
    
    def delete_document(self, document_id: str) -> bool:
        """Delete a document and its files"""
        doc_dir = self.upload_dir / document_id
        
        if not doc_dir.exists():
            return False
        
        try:
            # Delete all files in directory
            for file_path in doc_dir.iterdir():
                file_path.unlink()
            
            # Delete directory
            doc_dir.rmdir()
            
            logger.info(f"Deleted document: {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {e}")
            return False
    
    def get_statistics(self) -> Dict[str, Any]:
        """Get document manager statistics with performance metrics"""
        documents = self.list_documents()
        
        total_size = 0
        extensions = {}
        total_text_length = 0
        
        for doc in documents:
            ext = doc.get('extension', 'unknown')
            extensions[ext] = extensions.get(ext, 0) + 1
            
            try:
                total_size += int(doc.get('file_size', 0))
                total_text_length += int(doc.get('text_length', 0))
            except (ValueError, TypeError):
                pass
        
        avg_processing_time = 0
        if self.metrics["total_processed"] > 0:
            avg_processing_time = (
                self.metrics["total_processing_time"] / 
                self.metrics["total_processed"]
            )
        
        cache_hit_rate = 0
        total_requests = self.metrics["cache_hits"] + self.metrics["cache_misses"]
        if total_requests > 0:
            cache_hit_rate = self.metrics["cache_hits"] / total_requests
        
        return {
            'total_documents': len(documents),
            'total_size_bytes': total_size,
            'total_size_mb': round(total_size / (1024 * 1024), 2),
            'total_text_chars': total_text_length,
            'by_extension': extensions,
            'allowed_extensions': self.allowed_extensions,
            'performance': {
                'total_processed': self.metrics["total_processed"],
                'avg_processing_time_seconds': round(avg_processing_time, 2),
                'cache_hit_rate': round(cache_hit_rate, 2),
                'cache_hits': self.metrics["cache_hits"],
                'cache_misses': self.metrics["cache_misses"],
                'errors': self.metrics["errors"]
            },
            'engines': {
                'pymupdf_available': PYMUPDF_AVAILABLE,
                'pypdf2_available': PDF_AVAILABLE,
                'docx_available': DOCX_AVAILABLE
            }
        }
    
    def clear_cache(self) -> int:
        """Clear extraction cache"""
        count = 0
        
        try:
            for cache_file in self.cache_dir.iterdir():
                if cache_file.is_file() and cache_file.suffix == '.txt':
                    cache_file.unlink()
                    count += 1
            
            logger.info(f"Cleared {count} cached files")
            
        except Exception as e:
            logger.error(f"Error clearing cache: {e}")
        
        return count


# Global instance
_enhanced_doc_manager = None

def get_enhanced_document_manager() -> EnhancedDocumentManager:
    """Get or create enhanced document manager instance"""
    global _enhanced_doc_manager
    if _enhanced_doc_manager is None:
        _enhanced_doc_manager = EnhancedDocumentManager()
    return _enhanced_doc_manager
