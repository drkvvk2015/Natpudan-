"""
Document Management Service - Upload, process, and index medical documents
"""

import logging
import os
import hashlib
from typing import List, Dict, Any, Optional
from pathlib import Path
from datetime import datetime
import uuid

logger = logging.getLogger(__name__)

# Optional imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    logger.warning("PyPDF2 not available. Install with: pip install PyPDF2")
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available. Install with: pip install python-docx")
    DOCX_AVAILABLE = False


class DocumentManager:
    """
    Manages medical document uploads, processing, and indexing.
    Supports PDF, DOCX, and TXT files.
    """
    
    def __init__(
        self,
        upload_dir: str = "data/uploaded_documents",
        allowed_extensions: List[str] = None
    ):
        """
        Initialize document manager.
        
        Args:
            upload_dir: Directory to store uploaded files
            allowed_extensions: List of allowed file extensions
        """
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        if allowed_extensions is None:
            self.allowed_extensions = ['.pdf', '.docx', '.txt', '.md']
        else:
            self.allowed_extensions = allowed_extensions
        
        logger.info(f"Document manager initialized. Upload dir: {self.upload_dir}")
    
    def is_allowed_file(self, filename: str) -> bool:
        """
        Check if file extension is allowed.
        
        Args:
            filename: Filename to check
            
        Returns:
            True if allowed, False otherwise
        """
        return any(filename.lower().endswith(ext) for ext in self.allowed_extensions)
    
    async def save_upload(
        self,
        file_content: bytes,
        filename: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Save uploaded file and extract metadata.
        
        Args:
            file_content: File content bytes
            filename: Original filename
            metadata: Additional metadata
            
        Returns:
            Document information dictionary
        """
        if not self.is_allowed_file(filename):
            raise ValueError(f"File type not allowed: {filename}")
        
        # Generate unique document ID
        document_id = str(uuid.uuid4())
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        # Create document directory
        doc_dir = self.upload_dir / document_id
        doc_dir.mkdir(parents=True, exist_ok=True)
        
        # Save file
        file_path = doc_dir / filename
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        # Extract text content
        try:
            text_content = self.extract_text(file_path)
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            text_content = ""
        
        # Create document info
        doc_info = {
            'document_id': document_id,
            'filename': filename,
            'file_path': str(file_path),
            'file_size': len(file_content),
            'file_hash': file_hash,
            'extension': Path(filename).suffix.lower(),
            'uploaded_at': datetime.utcnow().isoformat(),
            'text_length': len(text_content),
            'metadata': metadata or {}
        }
        
        # Save metadata
        metadata_path = doc_dir / 'metadata.txt'
        with open(metadata_path, 'w', encoding='utf-8') as f:
            for key, value in doc_info.items():
                f.write(f"{key}: {value}\n")
        
        # Save extracted text
        text_path = doc_dir / 'extracted_text.txt'
        with open(text_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        logger.info(f"Saved document: {document_id} ({filename})")
        
        return doc_info
    
    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from various document formats.
        
        Args:
            file_path: Path to document
            
        Returns:
            Extracted text content
        """
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self._extract_pdf_text(file_path)
        elif extension == '.docx':
            return self._extract_docx_text(file_path)
        elif extension in ['.txt', '.md']:
            return self._extract_text_file(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        if not PDF_AVAILABLE:
            raise RuntimeError("PyPDF2 not installed")
        
        text = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text.append(f"\n--- Page {page_num + 1} ---\n")
                            text.append(page_text)
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num}: {e}")
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
            raise
        
        return '\n'.join(text)
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not installed")
        
        try:
            doc = DocxDocument(file_path)
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text.append(' | '.join(row_text))
            
            return '\n'.join(text)
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {e}")
            raise
    
    def _extract_text_file(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with latin-1 encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    def get_document(self, document_id: str) -> Optional[Dict[str, Any]]:
        """
        Get document information.
        
        Args:
            document_id: Document ID
            
        Returns:
            Document info or None if not found
        """
        doc_dir = self.upload_dir / document_id
        if not doc_dir.exists():
            return None
        
        # Read metadata
        metadata_path = doc_dir / 'metadata.txt'
        if not metadata_path.exists():
            return None
        
        doc_info = {}
        with open(metadata_path, 'r', encoding='utf-8') as f:
            for line in f:
                if ':' in line:
                    key, value = line.split(':', 1)
                    doc_info[key.strip()] = value.strip()
        
        return doc_info
    
    def get_document_text(self, document_id: str) -> Optional[str]:
        """
        Get extracted text for a document.
        
        Args:
            document_id: Document ID
            
        Returns:
            Extracted text or None if not found
        """
        text_path = self.upload_dir / document_id / 'extracted_text.txt'
        if not text_path.exists():
            return None
        
        with open(text_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def list_documents(self) -> List[Dict[str, Any]]:
        """
        List all uploaded documents.
        
        Returns:
            List of document information
        """
        documents = []
        
        for doc_dir in self.upload_dir.iterdir():
            if doc_dir.is_dir():
                doc_info = self.get_document(doc_dir.name)
                if doc_info:
                    documents.append(doc_info)
        
        # Sort by upload date (newest first)
        documents.sort(
            key=lambda x: x.get('uploaded_at', ''),
            reverse=True
        )
        
        return documents
    
    def delete_document(self, document_id: str) -> bool:
        """
        Delete a document and its files.
        
        Args:
            document_id: Document ID
            
        Returns:
            True if deleted, False if not found
        """
        doc_dir = self.upload_dir / document_id
        if not doc_dir.exists():
            return False
        
        try:
            # Delete all files in directory
            for file_path in doc_dir.iterdir():
                file_path.unlink()
            
            # Delete directory
            doc_dir.rmdir()
            
            logger.info(f"Deleted document: {document_id}")
            return True
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {e}")
            return False
    
    def get_statistics(self) -> Dict[str, Any]:
        """
        Get document manager statistics.
        
        Returns:
            Statistics dictionary
        """
        documents = self.list_documents()
        
        total_size = 0
        extensions = {}
        
        for doc in documents:
            # Count by extension
            ext = doc.get('extension', 'unknown')
            extensions[ext] = extensions.get(ext, 0) + 1
            
            # Sum file sizes
            try:
                total_size += int(doc.get('file_size', 0))
            except (ValueError, TypeError):
                pass
        
        return {
            'total_documents': len(documents),
            'total_size_bytes': total_size,
            'total_size_mb': round(total_size / (1024 * 1024), 2),
            'by_extension': extensions,
            'allowed_extensions': self.allowed_extensions
        }


# Global instance
_doc_manager = None

def get_document_manager() -> DocumentManager:
    """Get or create document manager instance"""
    global _doc_manager
    if _doc_manager is None:
        _doc_manager = DocumentManager()
    return _doc_manager
